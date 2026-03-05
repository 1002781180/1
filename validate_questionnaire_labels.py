"""
validate_questionnaire_labels.py
---------------------------------
验证 CFPS 2022 汇总问卷 DOCX 中的变量标签是否与分析代码中的变量映射一一对应。

运行方法：
    python validate_questionnaire_labels.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    import docx
except ImportError:
    sys.exit(
        "错误：未安装 python-docx。请先执行：pip install python-docx"
    )

# ===========================================================================
# 配置
# ===========================================================================
DOCX_FILE = "CFPS2022汇总问卷-241029done.docx"
OUTPUT_TEXT_FILE = "questionnaire_extracted_text.txt"
OUTPUT_REPORT_FILE = "label_validation_report.txt"

# ---------------------------------------------------------------------------
# 各脚本中定义的变量映射
# ---------------------------------------------------------------------------

# cfps_depression_analysis.py
DEPRESSION_FACTOR_MAP: dict[str, str] = {
    "年龄":     "age",
    "性别":     "gender",
    "城乡":     "urban22",
    "民族":     "minzu",
    "语文成绩": "wf501",
    "数学成绩": "wf502",
    "BMI":      "bmi",
    "近期生病": "wc0",
    "就医情况": "wc4_1",
    "慢性病诊断": "qp4001",
}
DEPRESSION_EMOTION_ITEMS: list[str] = [
    f"we3{str(i).zfill(2)}" for i in range(1, 13)
]
DEPRESSION_WEIGHT_CANDIDATES: list[str] = [
    "child_weight", "rswt_natcs22n", "rswt_natpn1022n"
]

# data_exploration.py
EXPLORATION_KEY_VARS: dict[str, list[str]] = {
    "慢性病诊断": ["qp4001", "wc4_1", "wc5ncode", "ill"],
    "抽样权重":   ["child_weight", "rswt_natcs22n", "rswt_natpn1022n"],
    "学业成绩":   ["wf501", "wf502"],
    "情绪行为(we3xx)": [f"we3{str(i).zfill(2)}" for i in range(1, 13)],
    "情绪行为(wn4xx)": ["wn401"],
}

# cfps_ml_pipeline.py
PIPELINE_COLUMN_MAP: dict[str, str] = {
    "age":            "age",
    "gender":         "gender",
    "urban":          "urban22",
    "sleep_duration": "wc2",
    "screen_time":    "wn401",
    "chronic_disease": "qp4001",
    "recent_illness": "wc0",
    "chinese_score":  "wf501",
    "math_score":     "wf502",
    "bmi":            "bmi",
    "weight":         "rswt_natcs22n",
}
PIPELINE_CHRONIC_FALLBACKS: list[str] = ["qp4001", "wc4_1", "wc0", "ill"]
PIPELINE_EMOTION_ITEMS: list[str] = [
    f"we3{str(i).zfill(2)}" for i in range(1, 13)
]

# ---------------------------------------------------------------------------
# 汇总：代码中所有出现的变量名 → 中文标签（尽可能收集）
# ---------------------------------------------------------------------------
def build_code_variable_map() -> dict[str, list[str]]:
    """
    返回 {cfps列名: [中文标签, ...]} 映射，聚合自三个脚本中的所有映射。
    同一变量可能在不同脚本中有不同中文标签，都保留在列表里。
    """
    # 反向映射：列名 → 中文标签集合
    mapping: dict[str, set[str]] = {}

    def _add(col: str, label: str) -> None:
        mapping.setdefault(col, set()).add(label)

    # cfps_depression_analysis.py
    for label, col in DEPRESSION_FACTOR_MAP.items():
        _add(col, label)
    for col in DEPRESSION_EMOTION_ITEMS:
        _add(col, "社会情绪发展条目")
    for col in DEPRESSION_WEIGHT_CANDIDATES:
        _add(col, "抽样权重")

    # data_exploration.py
    for label, cols in EXPLORATION_KEY_VARS.items():
        for col in cols:
            _add(col, label)

    # cfps_ml_pipeline.py
    pipeline_label_map = {
        "age":            "年龄",
        "gender":         "性别",
        "urban22":        "城乡",
        "wc2":            "睡眠时长（推测列名）",
        "wn401":          "屏幕时间/电视手机时长",
        "qp4001":         "慢性病诊断",
        "wc0":            "近期生病",
        "wf501":          "语文成绩",
        "wf502":          "数学成绩",
        "bmi":            "BMI",
        "rswt_natcs22n":  "抽样权重",
    }
    for col, label in pipeline_label_map.items():
        _add(col, label)
    for col in PIPELINE_CHRONIC_FALLBACKS:
        _add(col, "慢性病/健康状态候选")
    for col in PIPELINE_EMOTION_ITEMS:
        _add(col, "社会情绪发展条目")

    return {col: sorted(labels) for col, labels in mapping.items()}


# ===========================================================================
# DOCX 文本提取
# ===========================================================================
def extract_docx_text(path: str) -> str:
    """提取 DOCX 中的全部文本（段落 + 表格），返回合并字符串。"""
    doc = docx.Document(path)
    lines: list[str] = []

    # 段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                lines.append("\t".join(row_texts))

    return "\n".join(lines)


# ===========================================================================
# 变量名提取
# ===========================================================================
# CFPS 变量编号常见模式（按优先级，从长到短）
_VAR_PATTERNS: list[re.Pattern[str]] = [
    re.compile(r"\b(we3\d{2})\b", re.IGNORECASE),
    re.compile(r"\b(wf\d{3})\b", re.IGNORECASE),
    re.compile(r"\b(wc\d+(?:_\d+)?)\b", re.IGNORECASE),
    re.compile(r"\b(wn\d{3})\b", re.IGNORECASE),
    re.compile(r"\b(qp\d{4,})\b", re.IGNORECASE),
    re.compile(r"\b(rswt_\w+)\b", re.IGNORECASE),
    re.compile(r"\b(child_weight)\b", re.IGNORECASE),
    re.compile(r"\b(urban22)\b", re.IGNORECASE),
    re.compile(r"\b(minzu)\b", re.IGNORECASE),
    re.compile(r"\b(bmi)\b", re.IGNORECASE),
    re.compile(r"\b(age)\b", re.IGNORECASE),
    re.compile(r"\b(gender)\b", re.IGNORECASE),
    re.compile(r"\b(ill)\b", re.IGNORECASE),
]


def extract_variables_from_text(text: str) -> dict[str, list[str]]:
    """
    从文本中识别所有 CFPS 变量编号，并收集每个变量名周围的上下文（同行文字）
    作为"问卷描述"。

    返回 {变量名（小写）: [所有出现行的上下文列表]}
    """
    found: dict[str, list[str]] = {}
    for line in text.splitlines():
        for pattern in _VAR_PATTERNS:
            for match in pattern.finditer(line):
                var = match.group(1).lower()
                # 以整行作为上下文
                ctx = line.strip()
                found.setdefault(var, [])
                if ctx not in found[var]:
                    found[var].append(ctx)
    return found


# ===========================================================================
# 标签一致性检查
# ===========================================================================
def _label_consistent(code_labels: list[str], docx_contexts: list[str]) -> bool:
    """
    粗略检查代码中文标签是否在 DOCX 上下文中出现。
    只要任意一个标签出现在任意一个上下文中即视为一致。
    """
    for label in code_labels:
        for ctx in docx_contexts:
            if label in ctx:
                return True
    return False


# ===========================================================================
# 对比报告生成
# ===========================================================================
def build_report(
    code_vars: dict[str, list[str]],
    docx_vars: dict[str, list[str]],
) -> str:
    """生成对比报告文本，返回字符串。"""
    lines: list[str] = []
    sep = "=" * 70

    lines += [
        sep,
        "CFPS 2022 问卷变量标签验证报告",
        sep,
        "",
        "【说明】",
        "  ✅ 匹配：代码中使用，问卷中也出现，且中文标签能在问卷上下文中找到。",
        "  ⚠️  标签不一致：代码中使用，问卷中出现，但中文标签未在问卷上下文中找到。",
        "  ❌ 未找到：代码中使用，但问卷全文中未检索到该变量编号。",
        "  ❓ 潜在遗漏：问卷中出现，但代码中未使用的 CFPS 变量。",
        "",
    ]

    matched: list[str] = []
    label_mismatch: list[str] = []
    not_found: list[str] = []

    all_code_vars_lower = {v.lower(): (v, labels) for v, labels in code_vars.items()}

    for var_lower, (var_orig, code_labels) in sorted(all_code_vars_lower.items()):
        if var_lower in docx_vars:
            contexts = docx_vars[var_lower]
            if _label_consistent(code_labels, contexts):
                matched.append((var_orig, code_labels, contexts))
            else:
                label_mismatch.append((var_orig, code_labels, contexts))
        else:
            not_found.append((var_orig, code_labels))

    # 潜在遗漏：问卷中有但代码未使用
    code_var_set = set(all_code_vars_lower.keys())
    potentially_missed = {
        v: ctxs for v, ctxs in docx_vars.items()
        if v not in code_var_set
    }

    # ---------- ✅ 匹配 ----------
    lines.append(f"✅ 匹配的变量（共 {len(matched)} 个）")
    lines.append("-" * 50)
    if matched:
        for var, labels, contexts in matched:
            lines.append(f"  • {var}")
            lines.append(f"    代码标签：{', '.join(labels)}")
            # 只显示前 2 条上下文，避免过长
            for ctx in contexts[:2]:
                lines.append(f"    问卷上下文：{ctx[:120]}")
    else:
        lines.append("  （无）")
    lines.append("")

    # ---------- ⚠️ 标签不一致 ----------
    lines.append(f"⚠️  标签不一致的变量（共 {len(label_mismatch)} 个）")
    lines.append("-" * 50)
    if label_mismatch:
        for var, labels, contexts in label_mismatch:
            lines.append(f"  • {var}")
            lines.append(f"    代码标签：{', '.join(labels)}")
            for ctx in contexts[:2]:
                lines.append(f"    问卷上下文：{ctx[:120]}")
    else:
        lines.append("  （无）")
    lines.append("")

    # ---------- ❌ 未找到 ----------
    lines.append(f"❌ 代码中使用但问卷中未找到的变量（共 {len(not_found)} 个）")
    lines.append("-" * 50)
    if not_found:
        for var, labels in not_found:
            lines.append(f"  • {var}  （代码标签：{', '.join(labels)}）")
    else:
        lines.append("  （无）")
    lines.append("")

    # ---------- ❓ 潜在遗漏 ----------
    lines.append(f"❓ 问卷中出现但代码未使用的相关变量（共 {len(potentially_missed)} 个）")
    lines.append("-" * 50)
    if potentially_missed:
        for var in sorted(potentially_missed.keys()):
            ctxs = potentially_missed[var]
            lines.append(f"  • {var}")
            for ctx in ctxs[:1]:
                lines.append(f"    问卷上下文：{ctx[:120]}")
    else:
        lines.append("  （无）")
    lines.append("")

    # ---------- 汇总 ----------
    lines += [
        sep,
        "汇总统计",
        sep,
        f"  代码使用变量总数：{len(code_vars)}",
        f"  问卷中识别到的变量总数：{len(docx_vars)}",
        f"  ✅ 匹配：{len(matched)}",
        f"  ⚠️  标签不一致：{len(label_mismatch)}",
        f"  ❌ 代码有但问卷无：{len(not_found)}",
        f"  ❓ 问卷有但代码无：{len(potentially_missed)}",
        sep,
    ]

    return "\n".join(lines)


# ===========================================================================
# 主流程
# ===========================================================================
def main() -> None:
    docx_path = Path(DOCX_FILE)
    if not docx_path.exists():
        sys.exit(f"错误：未找到问卷文件 {DOCX_FILE}，请确保文件在当前目录下。")

    # 1. 提取 DOCX 文本
    print(f"正在读取问卷文件：{DOCX_FILE} …")
    full_text = extract_docx_text(str(docx_path))

    # 保存提取文本
    Path(OUTPUT_TEXT_FILE).write_text(full_text, encoding="utf-8")
    print(f"问卷全文已保存至：{OUTPUT_TEXT_FILE}（共 {len(full_text)} 字符）")

    # 2. 从 DOCX 提取变量名及上下文
    print("正在从问卷文本中提取变量编号 …")
    docx_vars = extract_variables_from_text(full_text)
    print(f"问卷中共识别到 {len(docx_vars)} 个不同变量编号：{sorted(docx_vars.keys())}")

    # 3. 从代码中汇总变量映射
    code_vars = build_code_variable_map()
    print(f"代码中共使用 {len(code_vars)} 个 CFPS 变量名")

    # 4. 生成对比报告
    report = build_report(code_vars, docx_vars)

    # 打印到控制台
    print()
    print(report)

    # 保存报告
    Path(OUTPUT_REPORT_FILE).write_text(report, encoding="utf-8")
    print(f"\n对比报告已保存至：{OUTPUT_REPORT_FILE}")


if __name__ == "__main__":
    main()
